"""Yanıt dışa aktarma — DOCX, XLSX, PDF."""

from __future__ import annotations

import io
import re
import textwrap
from typing import Optional

from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    content: str
    title:   Optional[str] = "Pimland AI Yaniti"
    format:  Optional[str] = None  # URL'den gelir, body'de opsiyonel


def _strip_md(text: str) -> str:
    """Markdown'dan düz metin — Word/PDF için."""
    text = re.sub(r"#{1,6}\s*", "", text)           # başlıklar
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)    # bold
    text = re.sub(r"\*(.+?)\*", r"\1", text)         # italic
    text = re.sub(r"`(.+?)`", r"\1", text)           # code
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)  # linkler
    text = re.sub(r"^[-*]\s+", "• ", text, flags=re.M)  # listeler
    return text.strip()


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _to_docx(content: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Başlık
    h = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = RGBColor(0xFF, 0x6B, 0x2B)

    # İçerik — satır satır işle
    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        # Heading tespiti
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            continue

        # Tablo satırı
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                continue  # ayırıcı satır
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = "Table Grid"
            for i, c in enumerate(cells):
                table.rows[0].cells[i].text = re.sub(r"\*\*(.+?)\*\*", r"\1", c)
            continue

        # Normal paragraf — bold/italic markup koru
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2]); run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1]); run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1]); run.font.name = "Courier New"
            else:
                p.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX ─────────────────────────────────────────────────────────────────────

def _to_xlsx(content: str, title: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Yanıt"

    # Başlık satırı
    ws.merge_cells("A1:D1")
    ws["A1"] = title
    ws["A1"].font = Font(size=14, bold=True, color="FF6B2B")
    ws["A1"].alignment = Alignment(horizontal="left")
    ws.row_dimensions[1].height = 24

    row = 3
    thin = Border(bottom=Side(style="thin", color="2A2A3A"))

    for line in content.split("\n"):
        line = line.rstrip()

        # Başlık
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            ws.merge_cells(f"A{row}:D{row}")
            cell = ws[f"A{row}"]
            cell.value = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2))
            cell.font = Font(bold=True, size=11 if len(m.group(1)) == 1 else 10)
            cell.fill = PatternFill("solid", fgColor="1A1A24")
            row += 1
            continue

        # Tablo satırı
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                continue
            for col_idx, val in enumerate(cells, 1):
                val_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", val)
                cell = ws.cell(row=row, column=col_idx, value=val_clean)
                if col_idx <= 4:
                    cell.alignment = Alignment(wrap_text=True)
            row += 1
            continue

        # Normal satır
        if line:
            clean = re.sub(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`", lambda m: m.group(1) or m.group(2) or m.group(3), line)
            ws.merge_cells(f"A{row}:D{row}")
            ws[f"A{row}"] = clean
            ws[f"A{row}"].alignment = Alignment(wrap_text=True)
        row += 1

    # Kolon genişlikleri
    ws.column_dimensions["A"].width = 60
    for col in "BCD":
        ws.column_dimensions[col].width = 20

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

def _to_pdf(content: str, title: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    # Türkçe karakter desteği için sistem fontu bul
    _FONT_NAME = "Helvetica"  # fallback
    for font_path in [
        "/System/Library/Fonts/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CustomFont", font_path))
                _FONT_NAME = "CustomFont"
            except Exception:
                pass
            break

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    orange = colors.HexColor("#FF6B2B")
    teal   = colors.HexColor("#00C2A8")
    light  = colors.HexColor("#F0F0F8")
    dark   = colors.HexColor("#13131A")

    title_style = ParagraphStyle("title", parent=styles["Title"],
        fontName=_FONT_NAME, textColor=orange, fontSize=16, spaceAfter=12)
    h1_style = ParagraphStyle("h1", parent=styles["Heading1"],
        fontName=_FONT_NAME, textColor=teal, fontSize=13, spaceBefore=10, spaceAfter=4)
    h2_style = ParagraphStyle("h2", parent=styles["Heading2"],
        fontName=_FONT_NAME, fontSize=11, spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle("body", parent=styles["Normal"],
        fontName=_FONT_NAME, fontSize=9, leading=14, spaceAfter=4)
    code_style = ParagraphStyle("code", parent=styles["Code"],
        fontSize=8, backColor=colors.HexColor("#1A1A24"),
        textColor=light, leading=12)

    story = [Paragraph(title, title_style), Spacer(1, 0.3*cm)]

    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            story.append(Spacer(1, 0.2*cm))
            continue

        # Başlık
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", m.group(2))
            st    = h1_style if level == 1 else h2_style
            story.append(Paragraph(text, st))
            continue

        # Tablo
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                continue
            row_data = [re.sub(r"\*\*(.+?)\*\*", r"\1", c) for c in cells]
            t = Table([row_data], colWidths=[4*cm] * min(len(row_data), 4))
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), dark),
                ("TEXTCOLOR",  (0,0), (-1,-1), light),
                ("FONTSIZE",   (0,0), (-1,-1), 8),
                ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#2A2A3A")),
                ("PADDING",    (0,0), (-1,-1), 4),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.1*cm))
            continue

        # Normal metin
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
        text = re.sub(r"\*(.+?)\*",     r"<i>\1</i>", text)
        text = re.sub(r"`(.+?)`",       r"<font name='Courier'>\1</font>", text)
        story.append(Paragraph(text, body_style))

    doc.build(story)
    return buf.getvalue()


# ── Endpoint ──────────────────────────────────────────────────────────────────

_MIME = {"docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
         "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
         "pdf":  "application/pdf"}

_BUILDERS = {"docx": _to_docx, "xlsx": _to_xlsx, "pdf": _to_pdf}


@router.post("/{fmt}")
async def export_response(fmt: str, req: ExportRequest):
    if fmt not in _BUILDERS:
        from fastapi import HTTPException
        raise HTTPException(400, f"Desteklenmeyen format: {fmt}. docx | xlsx | pdf")

    data  = _BUILDERS[fmt](req.content, req.title or "Yanıt")
    # Content-Disposition sadece ASCII kabul eder — Türkçe karakterleri temizle
    fname = re.sub(r"[^a-zA-Z0-9_\-]", "_", req.title or "yanit")[:40].strip("_") or "yanit"

    return StreamingResponse(
        io.BytesIO(data),
        media_type=_MIME[fmt],
        headers={"Content-Disposition": f"attachment; filename={fname}.{fmt}"},
    )
