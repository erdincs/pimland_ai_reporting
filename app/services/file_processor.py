"""Dosya içeriği çıkarıcı.

Desteklenen tipler: PDF, DOCX, XLSX, CSV, PNG, JPG
Her tip için ayrı işleyici; çıktı agent'a hazır format.
"""

from __future__ import annotations

import base64
import io
import os
import re
import time
from typing import Any, Dict, List

from app.core.logging import get_logger

log = get_logger(__name__)

MAX_FILE_BYTES  = 20 * 1024 * 1024   # 20 MB
MAX_IMG_BYTES   = 5  * 1024 * 1024   # 5 MB
MAX_TEXT_CHARS  = 50_000

ALLOWED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "text/csv",
    "text/plain",
    "image/png",
    "image/jpeg",
    "image/jpg",
}

ALLOWED_EXT = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".png", ".jpg", ".jpeg"}


# ── Güvenlik ──────────────────────────────────────────────────────────────────

def sanitize_filename(name: str) -> str:
    name = os.path.basename(name)
    name = re.sub(r"[^\w.\-]", "_", name)
    return name[:120]


def validate_file(filename: str, size: int, content_type: str) -> None:
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".xlsm":
        raise ValueError("Makro içerikli Excel dosyaları (.xlsm) kabul edilmez.")
    if ext not in ALLOWED_EXT:
        raise ValueError(f"Desteklenmeyen dosya tipi: {ext}")
    if content_type not in ALLOWED_MIME and not content_type.startswith("text/"):
        raise ValueError(f"Desteklenmeyen içerik tipi: {content_type}")
    if size > MAX_FILE_BYTES:
        raise ValueError(f"Dosya 20 MB sınırını aşıyor ({size // 1024 // 1024} MB).")


# ── PDF ───────────────────────────────────────────────────────────────────────

def process_pdf(data: bytes, filename: str) -> Dict[str, Any]:
    import fitz  # pymupdf
    doc   = fitz.open(stream=data, filetype="pdf")
    pages = doc.page_count
    parts = []
    for pg in doc:
        parts.append(pg.get_text())
    text = "\n".join(parts)
    truncated = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS]
    log.info("file.pdf_processed", filename=filename, pages=pages, chars=len(text))
    return {
        "type":      "document",
        "filename":  filename,
        "text":      text,
        "pages":     pages,
        "truncated": truncated,
    }


# ── DOCX ─────────────────────────────────────────────────────────────────────

def process_docx(data: bytes, filename: str) -> Dict[str, Any]:
    from docx import Document
    doc   = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    text = "\n".join(parts)
    truncated = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS]
    log.info("file.docx_processed", filename=filename, chars=len(text))
    return {
        "type":      "document",
        "filename":  filename,
        "text":      text,
        "truncated": truncated,
    }


# ── Excel / CSV ───────────────────────────────────────────────────────────────

def process_tabular(
    data: bytes,
    filename: str,
    session_id: str,
    db_url: str,
) -> Dict[str, Any]:
    import pandas as pd
    from sqlalchemy import create_engine, text as sa_text

    ext    = os.path.splitext(filename.lower())[1]
    ts     = int(time.time())
    tables = []

    if ext in (".xlsx", ".xls"):
        xl     = pd.ExcelFile(io.BytesIO(data))
        sheets = xl.sheet_names
    else:
        sheets = ["Sheet1"]

    sync_url = db_url.replace("+asyncpg", "+psycopg2")
    engine   = create_engine(sync_url, pool_pre_ping=True)

    for sheet in sheets:
        if ext in (".xlsx", ".xls"):
            df = pd.read_excel(io.BytesIO(data), sheet_name=sheet, nrows=10_000)
        else:
            df = pd.read_csv(io.BytesIO(data), nrows=10_000)

        safe_sheet = re.sub(r"\W+", "_", str(sheet).lower())[:20]
        safe_sid   = re.sub(r"\W+", "", session_id)[:12]
        pg_table   = f"tmp_{safe_sid}_{safe_sheet}_{ts}"

        df.to_sql(pg_table, engine, if_exists="replace", index=False)

        with engine.connect() as conn:
            conn.execute(sa_text(f"GRANT SELECT ON {pg_table} TO pimland_ro"))
            conn.commit()

        preview = df.head(5).to_csv(index=False)
        stats   = df.describe(include="all").to_string()

        tables.append({
            "sheet":    str(sheet),
            "pg_table": pg_table,
            "rows":     len(df),
            "columns":  list(df.columns),
            "preview":  preview,
            "stats":    stats[:2000],
        })
        log.info("file.tabular_loaded", sheet=sheet, pg_table=pg_table, rows=len(df))

    engine.dispose()
    return {
        "type":     "dataframe",
        "filename": filename,
        "tables":   tables,
    }


# ── Görsel ────────────────────────────────────────────────────────────────────

def process_image(data: bytes, filename: str, content_type: str) -> Dict[str, Any]:
    if len(data) > MAX_IMG_BYTES:
        raise ValueError(f"Görsel 5 MB sınırını aşıyor.")
    b64  = base64.standard_b64encode(data).decode("utf-8")
    mime = content_type if content_type.startswith("image/") else "image/jpeg"
    log.info("file.image_processed", filename=filename, bytes=len(data))
    return {
        "type":     "image",
        "filename": filename,
        "base64":   b64,
        "mime":     mime,
    }


# ── Ana dispatcher ────────────────────────────────────────────────────────────

def process_file(
    data: bytes,
    filename: str,
    content_type: str,
    session_id: str,
    db_url: str,
) -> Dict[str, Any]:
    filename = sanitize_filename(filename)
    validate_file(filename, len(data), content_type)

    ext = os.path.splitext(filename.lower())[1]

    if ext == ".pdf":
        return process_pdf(data, filename)
    if ext == ".docx":
        return process_docx(data, filename)
    if ext in (".xlsx", ".xls", ".csv"):
        return process_tabular(data, filename, session_id, db_url)
    if ext in (".png", ".jpg", ".jpeg"):
        return process_image(data, filename, content_type)
    if ext in (".txt",):
        text = data.decode("utf-8", errors="replace")[:MAX_TEXT_CHARS]
        return {"type": "document", "filename": filename, "text": text, "truncated": len(text) == MAX_TEXT_CHARS}

    raise ValueError(f"İşlenemeyen dosya tipi: {ext}")
